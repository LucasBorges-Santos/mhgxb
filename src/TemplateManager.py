import copy
import itertools
import json
import os
import random
import re
import string
import weakref

import matplotlib.colors
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import RGBColor
from docxtpl import DocxTemplate


class AncillaryDocx():
    def intervals_extract(self, iterable):
        iterable = sorted(set(iterable))
        for key, group in itertools.groupby(enumerate(iterable),
                                            lambda t: t[1] - t[0]):
            group = list(group)
            yield (group[0][1], group[-1][1]+1)

    def get_tags_from_text(self, str_html: str) -> list:
        """
            Return tags from a string:
                - "teste <a>teste</a>" -> ['teste ', <a>teste</a>]
        """
        tags_cords = []
        soup = BeautifulSoup(str_html, 'html.parser')

        for tag in soup.findAll(recursive=False):
            str_tag = str(tag)
            init_tag = str_html.find(str_tag)

            if init_tag == -1:
                raise Exception(f'Invalid Tag Format, check the spaces and tag format:\n    - {str_html}')
            
            finish_tag = str_html.find(str_tag) + len(str_tag)
            tags_cords.append((init_tag, finish_tag))

        if tags_cords:
            str_html_indexs = list(range(len(str_html)))

            for cord in tags_cords:
                for cord_index in list(range(cord[0], cord[1])):
                    str_html_indexs.pop(str_html_indexs.index(cord_index))
            tags_cords = list(self.intervals_extract(
                str_html_indexs)) + tags_cords
            sort_cords = sorted(tags_cords, key=lambda x: x[0])
            tags_splited = [str_html[init_tag:finish_tag]
                            for init_tag, finish_tag in sort_cords]
            return tags_splited
        else:
            return str_html

    def tag_to_paragraph_style(self, tag: BeautifulSoup, p) -> None:
        """
            Set style to each tag in text
        """
        if tag.find():
            element_style = list(tag.open_tag_counter)

            run = p.add_run(tag.text)

            if 'c' in element_style:
                color = tag.find('c').get('color')
                run.font.color.rgb = RGBColor.from_string(color)

            if 'b' in element_style:
                run.font.bold = True

            if 'i' in element_style:
                run.font.italic = True
        else:
            p.add_run(tag)

    def paragraph_style_to_tag_text(self, p) -> str:
        """
            get tag from paragraph
        """
        texts = ''
        for r in p.runs:
            text = r.text
            if r.font.color.rgb:
                text = f'<c color="{r.font.color.rgb}">{text}</c>'
            if r.font.bold:
                text = f'<b>{text}</b>'
            if r.font.italic:
                text = f'<i>{text}</i>'
            texts += text
        return texts


class TemplateDocx(AncillaryDocx):
    def __init__(self, template_path, df_tr: pd.DataFrame = False, savein: str = './', name: str = 'Default'):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.template_path = template_path
        self.extern_df_tr = False
        self.df_tr = self.get_translate_dataframe(df_tr)
        self.name = name
        self.savein = savein
        self.docx_path = os.path.join(self.savein, f'{self.name}.docx')
        self.xlsx_path = os.path.join(self.savein, f'{self.name}.xlsx')

    def get_paragraphs(self, doc, paragraphs=[]):
        """
        Return docx's text paragraphs.
        """
        for p in doc.paragraphs:
            if p.text:
                if p.text[0] == r'{':
                    continue
                if p.text.isspace():
                    continue
                paragraphs.append(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.get_paragraphs(cell, paragraphs)
        if not doc._parent:
            return paragraphs

    def render_docx(self, doc, df: pd.DataFrame, savein: str):
        """
        df - DataFrame with old and new text with this format:
            -    | OLD | NEW |
            - savein - save path with docx file name
        """

        for p in doc.paragraphs:
            if p.text:
                if p.text[0] == r'{':
                    continue
                if p.text.isspace():
                    continue

                if df.empty:
                    raise Exception("Render Recieved a Empty DataFrame!")
                new_text = df[df['OLD'] == p.text]['NEW'].iloc[0]
                new_text_splited = self.get_tags_from_text(new_text)

                p.text = ''
                for str_tag in new_text_splited:
                    tag = BeautifulSoup(str_tag, 'html.parser')
                    self.tag_to_paragraph_style(tag, p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.render_docx(cell, df, savein)
        if not doc._parent:
            doc.save(savein)

    def get_translate_dataframe(self, df_tr: pd.DataFrame = False):
        """
        Set Default df_tr to TemplateDocx Object \n
            - lang_str - Lang of final df_tr (if not df_tr) \n
            - df_tr - if inputed, set that as df_tr \n
        If without parameters:
            - Return a DataFrame with id of code and text for each paragraph

        """
        if isinstance(df_tr, pd.DataFrame):
            self.extern_df_tr = True
            return df_tr

        else:
            paragraphs = self.get_paragraphs(self.doc)
            for p in paragraphs:
                text = self.paragraph_style_to_tag_text(p)

            dict_temp = {'ID': [], 'LANG': []}
            for p in paragraphs:
                init_code = p.text.find('%[')
                final_code = p.text.find(']#')

                if init_code == -1:
                    continue

                id = p.text[init_code + 2:final_code]

                if id not in dict_temp['ID']:
                    text = p.text.replace(f'%[{str(id)}]#', '')
                    dict_temp['ID'].append(id)
                    dict_temp['LANG'].append(text)
            return pd.DataFrame(dict_temp)

    def gen_code_docx_df(self):
        """
        Generate a docx file, with your paragraphs with code:\n
            -   p.text + %[{str(id)}]#\n
        After, save df_dr as .xlsx file
        """
        code_docx = Document(self.template_path)
        paragraphs = self.get_paragraphs(code_docx)

        for p in paragraphs:
            if '%[' in p.text:
                raise Exception(
                    f"This file already have code, check the template_path.\n   - Paragraph with code: {p.text}")

        new_paragraphs = []
        new_paragraphs_without_code = []
        old_paragraphs = []
        indexs = []
        for index, p in enumerate(paragraphs):
            text_tag = self.paragraph_style_to_tag_text(p)
            if text_tag not in old_paragraphs:
                old_paragraphs.append(p.text)
                new_paragraphs_without_code.append(text_tag)

                text_tag = text_tag.replace("<", "<").replace(">", ">")
                new_paragraphs.append(f"{text_tag}%[{str(index)}]#")

                indexs.append(index)

        df = pd.DataFrame({
            'OLD': old_paragraphs,
            'NEW': new_paragraphs
        })
        df_tr = pd.DataFrame({
            'ID': indexs,
            'Default': new_paragraphs_without_code
        })

        df_tr.to_excel(self.xlsx_path)
        self.render_docx(code_docx, df, self.docx_path)

    def render_docx_translate(self, lang_out: str, savein: str) -> str:
        """
        Translate docx according its lang\n
            - lang_in: Document's lang\n
            - lang_out: output lang \n

        THE DOCX HAS TO HAVE CODE, SO MAKE THE FILE CODE USING "gen_code_docx_df".
        """
        docx_with_code = Document(self.template_path)
        paragraphs = self.get_paragraphs(docx_with_code)

        if lang_out not in self.df_tr.keys():
            raise Exception(
                "Lang out not found on Translate Dataframe, You can set df_tr manually or generate this columns using 'gen_code_docx_df' and translate the generate file.")

        has_code = False
        for p in paragraphs:
            if '%[' in p.text:
                has_code = True
        if not has_code:
            raise Exception(
                "Template Doesn't have Code! you can generate this template using 'gen_code_docx_df'.")

        new_paragraphs = []
        old_paragraphs = []

        for p in paragraphs:
            init_code = p.text.find('%[')
            final_code = p.text.find(']#')

            if init_code == -1:
                continue
            id_p = p.text[init_code + 2:final_code]
            if p not in old_paragraphs:
                new_text = self.df_tr[self.df_tr['ID'] == int(id_p)][lang_out]
                if new_text.empty:
                    continue
                old_paragraphs.append(p.text)
                new_paragraphs.append(
                    self.df_tr[self.df_tr['ID'] == int(id_p)][lang_out].iloc[0])

        df = pd.DataFrame({
            'OLD': old_paragraphs,
            'NEW': new_paragraphs
        })
        self.render_docx(self.doc, df, savein)
        return savein

    def __str__(self):
        info = f"Template Path: '{self.template_path}' Name: '{self.name}'\n" + \
            f"Savein: '{self.savein}' External DataFrame Translate: '{self.extern_df_tr}'\n" + \
            f"df_tr Columns Available: {self.df_tr.columns.values}"
        return info
